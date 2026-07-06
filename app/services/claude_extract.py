import base64
import json
import logging
from pathlib import Path

import anthropic

from app.config import settings

logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".csv"}

_SYSTEM_PROMPT_PATH = Path(__file__).parent.parent.parent / "system_prompt.txt"


def extract(file_bytes: bytes, filename: str) -> dict:
    system_prompt = _SYSTEM_PROMPT_PATH.read_text()

    client = anthropic.Anthropic(api_key=settings.anthropic_api_key)
    ext = _ext(filename)

    if ext == ".pdf":
        content = _pdf_content(file_bytes)
    else:
        text = _to_text(file_bytes, ext)
        content = [{"type": "text", "text": text}]

    logger.info("Calling Claude API for %s (%d bytes)", filename, len(file_bytes))
    with client.beta.messages.stream(
        model="claude-sonnet-4-6",
        max_tokens=64000,
        system=system_prompt,
        messages=[{"role": "user", "content": content}],
        betas=["output-128k-2025-02-19"],
    ) as stream:
        raw_text = stream.get_final_text()

    logger.info("Claude API response received for %s", filename)
    return _parse_json(raw_text)


def _ext(filename: str) -> str:
    return "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def _pdf_content(file_bytes: bytes) -> list:
    return [
        {
            "type": "document",
            "source": {
                "type": "base64",
                "media_type": "application/pdf",
                "data": base64.standard_b64encode(file_bytes).decode("utf-8"),
            },
        },
        {
            "type": "text",
            "text": "Extract all shipment data from this document.",
        },
    ]


def _to_text(file_bytes: bytes, ext: str) -> str:
    if ext == ".csv":
        return file_bytes.decode("utf-8", errors="replace")

    if ext == ".docx":
        from docx import Document
        import io
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts)

    if ext == ".xlsx":
        from openpyxl import load_workbook
        import io
        wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"[Sheet: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                if any(cell is not None for cell in row):
                    parts.append("\t".join("" if v is None else str(v) for v in row))
        return "\n".join(parts)

    raise ValueError(f"Unsupported file type: {ext}")


def _parse_json(raw_text: str) -> dict:
    text = raw_text.strip()
    if text.startswith("```"):
        lines = text.split("\n")
        text = "\n".join(l for l in lines if not l.strip().startswith("```")).strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        logger.error("JSON parse failed. Raw response (first 500 chars): %s", text[:500])
        raise
